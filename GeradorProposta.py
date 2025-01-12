# ================================================= MARTINS & SANTIAGO CORRETORA DE SEGUROS ================================================
# --------------------------------------------------- SEGURO SAÚDE | GERADOR DE PROPOSTAS --------------------------------------------------



## Propósito do Código
# Facilitar a geração de propostas de cotações de seguros de saúde para empresas PMEs

## Sumário
# 01. Intro
# 02. Cockpit
# 03. Functions
# 04. Componentes Streamlit
# 05. Pipeline Streamlit



# --------------------------------------------------------------- 01. INTRO ----------------------------------------------------------------



# Importa Bibliotecas
import os                                                                                                                                   # Biblioteca para manuseio de arquivos
import numpy as np                                                                                                                          # Biblioteca para manuseio de dados em matriz e distribuições
import pandas as pd                                                                                                                         # Biblioteca para manuseio de dados em DataFrame

from docx import Document
from docx.shared import Inches                                                                                                              # Import Inches from docx.shared

import time                                                                                                                                 # Biblioteca que permite realizar operações relacionadas com tempo 
from datetime import timedelta                                                                                                              # Biblioteca para calcular duração de trechos do código
from datetime import datetime                                                                                                               # Biblioteca para dizer data de hoje

from io import BytesIO                                                                                                                      # Biblioteca para manuseio de dados binários
import streamlit as st     

import matplotlib.pyplot as plt                                                                                                             # Import matplotlib.pyplot

# Limpa a memória
import gc
gc.collect()

# Clock inicio código
Start_Time = time.monotonic()

# Define o diretório onde está salvo os arquivos que serão utilizados
wdir = os.getcwd()                                                                                                                          # Guarda a localização do diretório do arquivo
wdir = wdir.replace("\\", "/")                                                                                                              # Troca o padrão de localização da Microsoft Windows para o padrão universal
os.chdir(wdir)                                                                                                                              # Define esse como o diretório padrão para esse algoritimo

# Cria um caminho para puxar os dados brutos e outro para o armazenamento dos resultados
Inputs_path  = "/01. Inputs/"
Results_path = "/02. Results/"

# Indica etapa do processo como concluida
print("01. Intro | OK")



# --------------------------------------------------------------- 02. COCKPIT --------------------------------------------------------------



# Hardinputs | Configurações
WORD_CONFIG_COLOR_HEADER_DEFAULT = "#585858"
LOGO_FULL_TEXT_PATH = "assets/M&S_Logo_Full_Text.png"
LOGO_PATH = "assets/M&S_Logo.png"



# -------------------------------------------------------------- 03. FUNCTIONS -------------------------------------------------------------



# Transformar o DataFrame em um arquivo Excel na memória
@st.cache_data
def convert_df_to_excel(
    df: pd.DataFrame
    ):
    """
    Função que converte um DataFrame em um arquivo Excel na memória

    Args:
        df (pd.DataFrame): DataFrame com dados

    Returns:
        output: Arquivo Excel na memória
    """
    # Cria um buffer BytesIO
    output = BytesIO()

    # Escreve o DataFrame no buffer BytesIO
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Sheet1')

    # Vai para o início do buffer BytesIO
    output.seek(0)

    return output


# Cria arquivo word
def create_word_file(
    WORD_CONFIG_LOGO_PATH: str = None
):
    """
    Função que cria um arquivo Word

    Args:
        None

    Returns:
        word_file: Arquivo Word na memória]
        doc: Documento Word
    """
    # Cria um arquivo Word
    word_file = BytesIO()
    doc = Document()
    doc.add_heading("Propostas de Seguro de Saúde", 0)

    # Caso o usuário não opte por adicionar o logo
    if WORD_CONFIG_LOGO_PATH == None:

        # Add a title or heading (optional)
        doc.add_heading("M&S Corretora de Seguros", level=1)

    # Caso o usuário opte por adicionar algum logo
    else:

        # Adiciona logo no arquivo Word
        doc.add_picture(WORD_CONFIG_LOGO_PATH, 
                        # width = Inches(1.25),
                        height = Inches(0.8))

    return word_file, doc


# Adiciona tabela no arquivo Word
def add_table_in_word(
    df: pd.DataFrame,
    company_name: str,
    WORD_CONFIG_WIDTH_TABLE: int,
    WORD_CONFIG_COLOR_HEADER: str,
    WORD_CONFIG_FONT_SIZE: int,
    WORD_CONFIG_FONT_COLOR: str,
    doc
):
    """
    Função que preparar um arquivo Word para receber as propostas

    Args:
        df (pd.DataFrame): DataFrame com as tabelas de cada empresa

    Returns:
        word_file: Arquivo Word na memória
    """
    # Renderiza a tabela no formato de imagem
    fig, ax = plt.subplots(figsize = (4, 1))
    ax.axis("tight")
    ax.axis("off")
    
    # Create the table
    table = ax.table(
        cellText = df.values,
        colLabels = df.columns,
        cellLoc = "center",
        loc = "center"
    )

    # Color the header
    for (i, j), cell in table.get_celld().items():
        if i == 0:  # Header row
            cell.set_facecolor(WORD_CONFIG_COLOR_HEADER)
            cell.set_fontsize(WORD_CONFIG_FONT_SIZE)
            cell.set_text_props(color = WORD_CONFIG_FONT_COLOR, weight = "bold")

    # Salva a imagem da tabela
    image_path = "temp/dataframe_image.png"
    plt.savefig(image_path, bbox_inches = "tight", dpi = 300)
    plt.close(fig)

    # Adiciona o nome da empresa no documento
    doc.add_heading(f"{company_name}", level = 1)

    # Adiciona a imagem da tabela no documento
    doc.add_picture(image_path, width = Inches(WORD_CONFIG_WIDTH_TABLE))

    return doc


# Função que coloca os dados no arquivo Word
def load_tables(
    Upload_Data: pd.DataFrame,
    WORD_CONFIG_LOGO_PATH: str = None
    ):
    """
    Função que coloca os dados no arquivo Word

    Args:
        Upload_Data (pd.DataFrame): DataFrame com dados

    Returns:
        word_file: Arquivo Word na memória
        doc: Documento Word
        file_name: Nome do arquivo
    """
    # Cria arquivo word
    word_file, doc = create_word_file(WORD_CONFIG_LOGO_PATH)

    # Pega a data atual
    now = datetime.now()
    time_now = now.strftime("%Y-%m-%d")

    # Pega o nome do cliente e da nome ao arquivo
    # client_name = Upload_Data["Config"].Cliente.values[0]
    file_name = f"Orçamento_{client_name}_{time_now}.docx"

    # Access each sheet by its name
    for sheet_name, df in Upload_Data.items():

        # Status
        print(f"Sheet name: {sheet_name}")
        print(df.head())

        # Adiciona tabela no arquivo Word
        doc = add_table_in_word(df, 
                                sheet_name, 
                                WORD_CONFIG_WIDTH_TABLE, 
                                WORD_CONFIG_COLOR_HEADER,
                                WORD_CONFIG_FONT_SIZE,
                                WORD_CONFIG_FONT_COLOR,
                                doc)

    # Use a BytesIO buffer to save the document in-memory
    word_file = BytesIO()
    doc.save(word_file)

    # Vai para o início do buffer BytesIO
    word_file.seek(0)  

    return word_file, doc, file_name



# ------------------------------------------------------- 04. COMPONENTES STREAMLIT --------------------------------------------------------


# Cria Sidebar
def sidebar(
    WORD_CONFIG_COLOR_HEADER_DEFAULT: str,
    LOGO_PATH: str
):
    """
    Função que cria a sidebar

    Args:
        None

    Returns:
        None
    """
    ################
    # Título Sidebar
    st.sidebar.header("M&S Corretora de Seguros")

    # Mostra o logo da empresa na sidebar
    st.sidebar.image(LOGO_PATH)

    ################
    # Configurações | Título
    expander = st.sidebar.expander("Configurações da Ferramenta")
    WORD_CONFIG_WIDTH_TABLE = expander.number_input(
        label = "Largura Tabelas", 
        value = 4, 
        help = "Escolhe a largura das tabelas que estarão no Word"
        )

    # Configurações | Cor Header
    WORD_CONFIG_COLOR_HEADER = expander.color_picker(
            label = "Cor do Header", 
            value = WORD_CONFIG_COLOR_HEADER_DEFAULT, 
            help = "Escolhe a cor do Header das tabelas que estarão no Word"
            )
    expander.write(WORD_CONFIG_COLOR_HEADER)

    # Configurações | Tamanho Fonte
    WORD_CONFIG_FONT_SIZE = expander.number_input(
        label = "Font Size Header", 
        value = 10,
        help = "Escolhe tamanho da fonte do Header das tabelas que estarão no Word"
        )

    # Configurações | Cor Fonte Header
    color_options_dict = {"Branco": "#FFFFFF", 
                          "Preto": "#000000"}
    color_options_list = color_options_dict.keys()
    color_chosen = expander.selectbox(
        "Cor Fonte Header", 
        color_options_list, 
        help = "Escolhe a cor da fonte do Header das tabelas que estarão no Word"
        )
    WORD_CONFIG_FONT_COLOR = color_options_dict[color_chosen]

    # Configurações | Logo
    logo_options_dict = {"Logo Completo": LOGO_FULL_TEXT_PATH,
                         "Logo Simples": LOGO_PATH,
                         "Sem Logo": None
                        }
    logo_options_list = logo_options_dict.keys()
    logo_chosen = expander.selectbox(
        "Tipo de Logo", 
        logo_options_list,
        help = "Escolhe formato do logo que estará no Header do Word",
        )
    WORD_CONFIG_LOGO_PATH = logo_options_dict[logo_chosen]

    ################
    # Passo-a-Passo
    expander = st.sidebar.expander(label = "Passo-a-Passo")
    expander.write('''
        1) Faça o upload do arquivo Excel com os dados da proposta
        2) Clique no botão "Gerador de Proposta"
        3) Quando ficar pronto um botão de download do word aparecerá
    ''')

    return WORD_CONFIG_FONT_SIZE, WORD_CONFIG_WIDTH_TABLE, WORD_CONFIG_COLOR_HEADER, WORD_CONFIG_FONT_COLOR, WORD_CONFIG_LOGO_PATH

# Botão de Upload de Arquivo
def button_upload_file():
    """
    Função que cria o botão de upload de arquivo

    Args:
        None

    Returns:
        df_result: DataFrame com dados
    """
    # Permite usuário dar upload de arquivo com dados no formato para construir os gráficos
    Upload_Data = st.file_uploader(
        label = "Upload Proposta Excel", 
        type = ["xlsx"],
        help = "Arraste o arquivo Excel com as tabelas coladas nas abas e às nomeie com o nome da seguradora"
        )

    # Caso tenha sido realizado algum upload
    if Upload_Data is not None:

        # Pega valor de bytes do arquivo
        bytes_data = Upload_Data.getvalue()

        # Le arquivo excel
        df_result = pd.read_excel(BytesIO(bytes_data), sheet_name = None)  

    # Caso não haja upload, usa dados de um exemplo
    if Upload_Data is None:

        # Le arquivo com dados template
        df_result = pd.read_excel("Exemplo_Proposta_Saude.xlsx", engine = "openpyxl", sheet_name = None)

    return df_result

# Botão Gerador de Propostas
def gerador_propostas(
    df_result: pd.DataFrame,
    WORD_CONFIG_LOGO_PATH: str = None,
    client_name: str = "Nome_do_Cliente"
):
    """
    Função que gera propostas de seguro de saúde

    Args:
        df_result (pd.DataFrame): DataFrame com dados

    Returns:
        None
    """
    # Carrega as tabelas de cotação e cola no Word
    word_file, doc, file_name = load_tables(df_result, WORD_CONFIG_LOGO_PATH)

    # Cria botão para download do arquivo Word
    st.download_button(
        label = f"Download Proposta {client_name}", 
        data = word_file, 
        file_name = file_name
        )

    return None


# --------------------------------------------------------- 05. PIPELINE STREAMLIT ---------------------------------------------------------



# Título
st.title("Propostas de Seguro de Saúde")

# Input string nome do cliente
st.write("Insira o nome do cliente")
client_name = st.text_input("Nome do Cliente", "Nome_Exemplo")

# Cria a sidebar
WORD_CONFIG_FONT_SIZE, WORD_CONFIG_WIDTH_TABLE, WORD_CONFIG_COLOR_HEADER, WORD_CONFIG_FONT_COLOR, WORD_CONFIG_LOGO_PATH = sidebar(WORD_CONFIG_COLOR_HEADER_DEFAULT, LOGO_PATH)

# Botão Upload
df_result = button_upload_file()

########## Geador de Proposta ##########
# Caso botão de gerar proposta seja clicado
if st.button("Gerador de Proposta", help = "Ao clicar, aguarde alguns segundos enquanto o arquivo é gerado e depois clique no botão de download para obter o arquivo Word"):

    gerador_propostas(df_result, WORD_CONFIG_LOGO_PATH, client_name)